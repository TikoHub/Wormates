import os
import re
import traceback

import magic
from django.shortcuts import redirect
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from rest_framework.decorators import api_view, permission_classes
from rest_framework.pagination import PageNumberPagination, LimitOffsetPagination
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.permissions import IsAuthenticated, IsAuthenticatedOrReadOnly, AllowAny
from .models import CommentLike, CommentDislike, BookUpvote, BookDownvote
from django.contrib.auth.mixins import LoginRequiredMixin
from .forms import BooksForm, SeriesForm
from .filters import BookFilter
from django.utils.decorators import method_decorator
from django.contrib.admin.views.decorators import staff_member_required
from django.http import HttpResponse
from django.db.models import Max, F
from users.models import Notification, Wallet
from rest_framework.views import APIView
from rest_framework import status, permissions
from rest_framework import generics
from rest_framework.response import Response
from .models import Vote, UserBookHistory
from .utils import get_client_ip, is_book_purchased_by_user, log_book_access
from .serializer import *
from .converters import create_fb2, parse_fb2
from django.core.files.storage import default_storage
from datetime import date
from users.models import WebPageSettings, Library, Profile, FollowersCount, PurchasedBook, NotificationSettings, \
    UserMainPageSettings
from django.db.models import Exists, OuterRef
from django.db.models import Q
from rest_framework.exceptions import PermissionDenied
import logging
from django.http import FileResponse, Http404
from .permissions import IsBookAuthor
from users.notification_utils import send_book_update_notifications
from lxml import etree
from docx import Document
import fitz
import zipfile
from bs4 import BeautifulSoup
from django.db import transaction
from channels.layers import get_channel_layer
from asgiref.sync import async_to_sync

from users.models import UserBookChapterNotification


class InfiniteScrollPagination(LimitOffsetPagination):
    default_limit = 48
    max_limit = 96


class BooksListAPIView(generics.ListAPIView):
    serializer_class = MainPageSerializer
    pagination_class = InfiniteScrollPagination

    def get_queryset(self):
        user = self.request.user
        genre_id = self.request.query_params.get('genre_id', None)
        date_filter = self.request.query_params.get('date_filter', None)

        base_query = Book.objects.exclude(genre__name='Undefined')

        # Фильтрация по жанру с учётом дочерних жанров
        if genre_id:
            try:
                selected_genre = Genre.objects.get(id=genre_id)
                # Получаем все дочерние жанры
                all_genres = {selected_genre}
                all_genres.update(self.get_all_subgenres_iterative(selected_genre))
                # Получаем список ID всех жанров
                genre_ids = [genre.id for genre in all_genres]
                base_query = base_query.filter(genre__id__in=genre_ids)
            except Genre.DoesNotExist:
                pass  # Если жанр не найден, можно проигнорировать фильтр или вернуть пустой queryset

        # Остальная часть вашего кода остается без изменений...

        # Фильтрация по дате выпуска
        if date_filter:
            # Ваш существующий код фильтрации по дате
            pass

        # Получение настроек пользователя
        if user.is_authenticated:
            try:
                settings = user.main_page_settings
            except UserMainPageSettings.DoesNotExist:
                settings = None
        else:
            settings = None

        # Применение настроек пользователя к запросу
        if settings:
            if settings.show_first_books:
                base_query = base_query.filter(volume_number=1)
            if settings.show_only_free_books:
                base_query = base_query.filter(price=0)
            if settings.restricted_mode:
                base_query = base_query.filter(is_adult=False)
        else:
            # Настройки по умолчанию для неаутентифицированных пользователей
            base_query = base_query.filter(is_adult=False)

        # Проверка наличия опубликованных глав
        has_published_chapters = Chapter.objects.filter(
            book=OuterRef('pk'),
            published=True
        )
        base_query = base_query.annotate(
            qualified=Exists(has_published_chapters)
        )

        # Фильтрация по видимости
        if user.is_authenticated:
            base_query = base_query.filter(
                Q(visibility='public') |
                Q(visibility='followers', author__follower_users__follower=user),
                qualified=True
            ).distinct()
        else:
            base_query = base_query.filter(
                visibility='public',
                qualified=True
            )

        return base_query

    def get_all_subgenres_iterative(self, genre):
        subgenres = set()
        queue = [genre]
        while queue:
            current_genre = queue.pop(0)
            child_genres = current_genre.child_subgenres.all()
            for child in child_genres:
                if child not in subgenres:
                    subgenres.add(child)
                    queue.append(child)
        return subgenres


class TopViewedBooksAPIView(BooksListAPIView):
    pagination_class = InfiniteScrollPagination

    def get_viewed_queryset(self):
        base_query = self.get_queryset()
        return base_query.order_by('-views_count')


class TopRatedBooksAPIView(BooksListAPIView):
    pagination_class = InfiniteScrollPagination

    def get_rated_queryset(self):
        base_query = self.get_queryset()
        return base_query.order_by('-rating')


class BookDetailAPIView(generics.RetrieveAPIView):
    queryset = Book.objects.all()
    serializer_class = BookSerializer

    def get_object(self):
        book_id = self.kwargs.get('book_id')
        book = get_object_or_404(Book, id=book_id)

        # Allow access to unlisted books only via direct link
        if book.visibility == 'unlisted':
            return book

        # Check if the user has access to the book based on its visibility
        user = self.request.user
        if book.visibility == 'private' and book.author != user:
            raise PermissionDenied('You do not have permission to view this book.')
        elif book.visibility == 'followers' and not book.author.followers.filter(user=user).exists():
            raise PermissionDenied('You do not have permission to view this book.')

        if book.is_adult:
            if not user.is_authenticated:
                raise Http404("This Book is for Adults Only. Please Sign in to Continue.")
            if hasattr(user.profile, 'webpagesettings') and user.profile.webpagesettings.date_of_birth:
                date_of_birth = user.profile.webpagesettings.date_of_birth
                age = (date.today() - date_of_birth).days // 365
                if age < 18:
                    raise Http404("This Book is for Adults Only.")
            else:
                raise Http404("Your age is not specified.")

        return book

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        user = request.user
        should_increment_view = False

        # Логика увеличения счётчика просмотров
        if user.is_authenticated:
            # Для аутентифицированных пользователей
            last_view, created = BookView.objects.get_or_create(
                user=user,
                book=instance,
                defaults={'last_viewed': timezone.now()}
            )
            if not created:
                time_since_last_view = timezone.now() - last_view.last_viewed
                if time_since_last_view > timedelta(hours=24):
                    should_increment_view = True
                    last_view.last_viewed = timezone.now()
                    last_view.save()
            else:
                should_increment_view = True
        else:
            # Для неаутентифицированных пользователей
            ip_address = get_client_ip(request)
            if ip_address:
                last_view, created = BookView.objects.get_or_create(
                    ip_address=ip_address,
                    book=instance,
                    defaults={'last_viewed': timezone.now()}
                )
                if not created:
                    time_since_last_view = timezone.now() - last_view.last_viewed
                    if time_since_last_view > timedelta(hours=12):
                        should_increment_view = True
                        last_view.last_viewed = timezone.now()
                        last_view.save()
                else:
                    should_increment_view = True
            else:
                # Не удалось получить IP-адрес
                pass

        if should_increment_view:
            # Атомарно увеличиваем счётчик просмотров
            Book.objects.filter(id=instance.id).update(views_count=F('views_count') + 1)
            # Обновляем объект instance
            instance.refresh_from_db()

        serialized_data = self.get_serializer(instance).data

        # Ваш существующий код по обновлению истории чтения
        if user.is_authenticated:
            if user.profile.record_history:
                UserBookHistory.objects.update_or_create(
                    user=user,
                    book=instance,
                    defaults={'last_accessed': timezone.now()}
                )
        else:
            # Логика для неаутентифицированных пользователей
            unlogged_user_history = request.session.get('unlogged_user_history', [])
            if instance.id not in unlogged_user_history:
                unlogged_user_history.append(instance.id)
                if len(unlogged_user_history) > 10:
                    unlogged_user_history.pop(0)
                request.session['unlogged_user_history'] = unlogged_user_history

        return Response(serialized_data)


class BookDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, book_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)

        # Проверяем, что пользователь является автором книги
        if book.author != user:
            return Response({'error': 'You are not authorized to delete this book.'},
                            status=status.HTTP_403_FORBIDDEN)

        confirm_delete = request.query_params.get('confirm_delete', 'false').lower()
        if confirm_delete == 'true':
            # Удаляем книгу и все связанные главы
            book.delete()
            return Response({'message': 'Book deleted successfully.'}, status=status.HTTP_200_OK)
        else:
            # Возвращаем предупреждение
            return Response({
                'warning': 'You are about to delete the book, the changes are irreversible. Do you want to proceed?',
                'confirmation_required': True
            }, status=status.HTTP_200_OK)


@api_view(['GET'])
def get_book_info(request, book_id):
    try:
        book = Book.objects.get(pk=book_id)
    except Book.DoesNotExist:
        return Response({'error': 'Book not found'}, status=404)

    serializer = BookInfoSerializer(book)
    return Response(serializer.data)


@api_view(['GET'])
def get_book_content(request, book_id):
    try:
        book = Book.objects.get(pk=book_id)
    except Book.DoesNotExist:
        return Response({'error': 'Book not found'}, status=404)

    serializer = BookContentSerializer(book)
    return Response(serializer.data)


class GenreListAPIView(generics.ListAPIView):
    queryset = Genre.objects.all()
    serializer_class = GenreSerializer


class BookSearch(ListView):
    template_name = 'store/book_search.html'
    queryset = Book.objects.all()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filter'] = BookFilter(self.request.GET, queryset=self.get_queryset())
        return context


class SimilarBooksAPIView(APIView):
    permission_classes = [AllowAny]

    def get(self, request, book_id):
        book = get_object_or_404(Book, id=book_id)
        main_genre = book.genre
        subgenres = set(book.subgenres.all())

        recommended = []

        # Проверка на следующую/предыдущую книгу из серии
        if book.series:
            next_book = Book.objects.filter(series=book.series, volume_number=book.volume_number + 1).first()
            if next_book:
                recommended.append(next_book)
            else:
                if book.volume_number > 1:
                    prev_book = Book.objects.filter(series=book.series, volume_number=book.volume_number - 1).first()
                    if prev_book:
                        recommended.append(prev_book)

        # Основной набор - книги того же жанра, исключая текущую и книги из серии (кроме уже добавленных)
        base_queryset = Book.objects.filter(genre=main_genre).exclude(id=book.id)
        if book.series:
            base_queryset = base_queryset.exclude(series=book.series)

        # Предварительно подгружаем subgenres и views_count (если views_count у вас отдельным образом)
        other_books = list(base_queryset.prefetch_related('subgenres'))

        book_scores = []
        for candidate in other_books:
            candidate_subgenres = set(candidate.subgenres.all())
            intersection_count = len(subgenres.intersection(candidate_subgenres))

            # Рассчёт очков
            # 10 очков за совпадение главного жанра (у нас уже гарантированно совпадает, т.к. filter по genre)
            points = 10 + (2 * intersection_count)

            # Сохраняем (книга, очки)
            book_scores.append((candidate, points))

        # Сортируем по points убыванию, при равенстве - по views_count убыванию
        # Предполагается, что у модели Book есть поле views_count
        # Если его нет, используйте другое поле, или не сортируйте по нему.
        book_scores.sort(key=lambda x: (x[1], x[0].views_count), reverse=True)

        max_recs = 5
        already_added = len(recommended)
        needed = max_recs - already_added

        # Берем top 'needed' книг из book_scores
        chosen = [item[0] for item in book_scores[:needed]]

        final_recommendations = recommended + chosen

        # Проверяем, набрали ли мы 5 рекомендаций
        # Если меньше 5 (например, книг в этом жанре мало),
        # можно сделать fallback: взять оставшиеся книги (даже с 0 очков), или не брать.
        # В данном примере, если мало книг по жанру, просто берем что есть.
        # Если хотим строго 5, то можем дополнить другими книгами (не из того же жанра)
        # Но этого вы не просили. Если надо - можно реализовать.

        if len(final_recommendations) < 5:
            # Если хотим строго 5, добавим любые книги из базы, исключая текущую
            # и из серии, если возможно. Это fallback
            fallback_needed = 5 - len(final_recommendations)
            fallback_qs = Book.objects.exclude(id=book.id)
            if book.series:
                fallback_qs = fallback_qs.exclude(series=book.series)
            # Исключаем уже выбранные
            fallback_ids = [b.id for b in final_recommendations]
            fallback_qs = fallback_qs.exclude(id__in=fallback_ids)
            fallback_books = list(fallback_qs[:fallback_needed])
            final_recommendations += fallback_books

        serializer = ReccomendationSerializer(final_recommendations, many=True, context={'request': request})
        return Response(serializer.data)


class StudioWelcomeAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        serializer = BookTypeSerializer(data=request.data)
        if serializer.is_valid():
            book_type = serializer.validated_data.get('book_type')

            with transaction.atomic():
                new_book = Book.objects.create(author=request.user, book_type=book_type)

                # Создаем первую главу
                first_chapter = Chapter.objects.create(
                    book=new_book,
                    chapter_number=1
                )

                return Response({
                    'message': 'Book created successfully',
                    'book_id': new_book.id,
                    'first_chapter_id': first_chapter.id,
                }, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class StudioBooksAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        books = Book.objects.filter(author=request.user)
        serializer = StudioBookSerializer(books, context={'request': request}, many=True)
        return Response(serializer.data)

    def patch(self, request, book_id):
        book = Book.objects.get(id=book_id, author=request.user)
        serializer = BookVisibilitySerializer(book, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class StudioSeriesAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Get series created by the user
        series = Series.objects.filter(author=request.user).prefetch_related('books')
        series_data = []
        for serie in series:
            books = serie.books.all()
            books_data = StudioSeriesBooksSerializer(books, many=True, context={'request': request}).data
            series_data.append({
                'id': serie.id,
                'name': serie.name,
                'books': books_data
            })

        # Get books that are not part of any series
        standalone_books = Book.objects.filter(author=request.user, series__isnull=True)
        standalone_books_data = StudioSeriesBooksSerializer(standalone_books, context={'request': request}, many=True).data

        # Combine series and standalone books in the response
        response_data = {
            'series': series_data,
            'standalone_books': standalone_books_data
        }
        return Response(response_data)

    def put(self, request, book_id):
        book = get_object_or_404(Book, id=book_id, author=request.user)
        old_series = book.series
        serializer = StudioBookSerializer(book, data=request.data, partial=True)
        if serializer.is_valid():
            # Проверяем, передана ли новая серия в запросе
            if 'series' in serializer.validated_data:
                new_series = serializer.validated_data['series']
                if new_series is not None:
                    # Проверяем, принадлежит ли серия текущему пользователю
                    if new_series.author != request.user:
                        return Response(
                            {"error": "You do not have permission to add this book to the selected series."},
                            status=status.HTTP_403_FORBIDDEN
                        )
            # Если проверка пройдена - сохраняем изменения
            serializer.save()

            # Если изменились 'volume_number' или 'series', вызываем обновление номеров томов
            if 'volume_number' in serializer.validated_data or 'series' in serializer.validated_data:
                self.update_volume_numbers(book, old_series=old_series)
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def update_volume_numbers(self, updated_book, old_series=None):
        # Update volume numbers in the old series if the book was moved from another series
        if old_series:
            old_series_books = old_series.books.exclude(id=updated_book.id).order_by('volume_number')
            for i, book in enumerate(old_series_books, start=1):
                book.volume_number = i
                book.save(update_fields=['volume_number'])

        # Update volume numbers in the new series if the book is part of a series
        if updated_book.series:
            new_series_books = updated_book.series.books.exclude(id=updated_book.id).order_by('volume_number')
            new_books_list = list(new_series_books)
            new_books_list.insert(updated_book.volume_number - 1, updated_book)
            for i, book in enumerate(new_books_list, start=1):
                book.volume_number = i
                book.save(update_fields=['volume_number'])


class StudioCommentsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # Get only top-level comments for books where the logged-in user is the author
        comments = Comment.objects.filter(book__author=request.user, parent_comment__isnull=True).order_by('-timestamp')
        serializer = StudioCommentSerializer(comments, many=True, context={'request': request})
        return Response(serializer.data)

    def post(self, request, book_id):
        book = get_object_or_404(Book, pk=book_id)
        if not book.can_user_comment(request.user):
            return Response({'error': 'You are not allowed to comment on this book.'}, status=status.HTTP_403_FORBIDDEN)

        data = request.data.copy()  # Make a mutable copy of the request data
        data['book'] = book.id  # Ensure the book is correctly associated

        serializer = CreateCommentSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            new_comment = serializer.save(user=request.user)  # User is added here
            return Response(StudioCommentSerializer(new_comment, context={'request': request}).data,
                            status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, comment_id):
        comment = get_object_or_404(Comment, id=comment_id)
        if comment.book.author != request.user:
            return Response({'error': 'You are not authorized to delete this comment.'},
                            status=status.HTTP_403_FORBIDDEN)

        comment.delete()
        return Response({'status': 'Comment deleted successfully'}, status=status.HTTP_204_NO_CONTENT)


class StudioIllustrationsAPIView(APIView):
    permission_classes = [IsAuthenticated, IsBookAuthor]
    parser_classes = (MultiPartParser, FormParser)

    def get(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        try:
            book = Book.objects.get(id=book_id)
        except Book.DoesNotExist:
            return Response({'error': 'Book not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the request user is the author of the book
        self.check_object_permissions(request, book)

        book_serializer = BookSerializer(book, context={'request': request})
        illustrations = Illustration.objects.filter(book=book)
        illustrations_serializer = IllustrationSerializer(illustrations, many=True, context={'request': request})

        return Response({
            'cover_page': book_serializer.data['coverpage'],
            'illustrations': illustrations_serializer.data
        })

    def post(self, request, *args, **kwargs):
        print("request.FILES:", request.FILES)
        print("request.data:", request.data)

        book_id = kwargs.get('book_id')
        try:
            book = Book.objects.get(id=book_id)
        except Book.DoesNotExist:
            return Response({'error': 'Book not found'}, status=status.HTTP_404_NOT_FOUND)

        # Set the book's cover page
        cover_image = request.FILES.get('cover_image')
        if cover_image:
            book.coverpage = cover_image
            book.save()

        # Add illustrations with descriptions
        for key in request.FILES.keys():
            if key.startswith('illustration_image_'):
                index = key.split('_')[-1]
                image = request.FILES[key]
                description = request.data.get(f'illustration_description_{index}', '')

                serializer = IllustrationSerializer(data={'image': image, 'description': description},
                                                    context={'request': request})
                if serializer.is_valid():
                    serializer.save(book=book)
                else:
                    print(serializer.errors)  # Добавьте эту строку для отладки
                    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        return Response({'message': 'Cover page and illustrations updated successfully'},
                        status=status.HTTP_201_CREATED)

    def put(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        illustration_id = kwargs.get('illustration_id')

        try:
            illustration = Illustration.objects.get(id=illustration_id, book__id=book_id)
        except Illustration.DoesNotExist:
            return Response({'error': 'Illustration not found'}, status=status.HTTP_404_NOT_FOUND)

        serializer = IllustrationSerializer(illustration, data=request.data, partial=True, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        illustration_id = request.data.get('illustration_id')
        use_for_library_cover = request.data.get('use_for_library_cover', False)

        try:
            book = Book.objects.get(id=book_id)
            illustration = Illustration.objects.get(id=illustration_id, book=book)
        except (Book.DoesNotExist, Illustration.DoesNotExist):
            return Response({'error': 'Book or Illustration not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the request user is the author of the book
        self.check_object_permissions(request, book)

        illustration.use_for_library_cover = use_for_library_cover
        illustration.save()

        return Response({'message': 'Illustration updated successfully'}, status=status.HTTP_200_OK)

    def delete(self, request, *args, **kwargs):
        book_id = kwargs.get('book_id')
        illustration_id = kwargs.get('illustration_id')

        try:
            illustration = Illustration.objects.get(id=illustration_id, book__id=book_id)
        except Illustration.DoesNotExist:
            return Response({'error': 'Illustration not found'}, status=status.HTTP_404_NOT_FOUND)

        illustration.delete()
        return Response({'message': 'Illustration deleted successfully'}, status=status.HTTP_204_NO_CONTENT)


class BooksCreateAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        serializer = BookCreateSerializer(data=request.data, context={'request': request})

        if serializer.is_valid():
            book = serializer.save()
            request.session['book_id'] = book.id
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ChapterContentView(APIView):
    def post(self, request, *args, **kwargs):
        serializer = ChapterContentSerializer(data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ChapterListView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id):
        book = Book.objects.get(id=book_id)
        chapters = book.chapters.all().order_by('created')  # Assuming you want them ordered by creation date
        serializer = ChapterSideSerializer(chapters, many=True)
        return Response(serializer.data)

    def post(self, request, book_id):
        serializer = ChapterSerializers(data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save(book_id=book_id)  # Ensure the book_id is passed correctly to the save method
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class AddChapterView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        try:
            with transaction.atomic():
                book = Book.objects.select_for_update().get(id=book_id)  # Lock the book record for update
                if book.author != request.user:
                    return Response({'error': 'You are not the author of this book.'}, status=status.HTTP_403_FORBIDDEN)

                # Refresh the book instance to ensure it's up-to-date
                book.refresh_from_db()

                # Check if the book is of type "Short Story / Poem" and already has a chapter
                if book.book_type == 'short_story_poem' and book.chapters.exists():
                    return Response({'error': 'Short Story / Poem books can only have one chapter.'},
                                    status=status.HTTP_400_BAD_REQUEST)

                # Create an empty chapter
                last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0
                next_chapter_number = last_chapter_number + 1

                # Create an empty chapter with the next chapter number
                new_chapter = Chapter.objects.create(book=book, chapter_number=next_chapter_number)
                serializer = ChapterSerializers(new_chapter)

                return Response(serializer.data, status=status.HTTP_201_CREATED)

        except Book.DoesNotExist:
            return Response({'error': 'Book not found.'}, status=status.HTTP_404_NOT_FOUND)


class StudioChapterView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, chapter_id, format=None):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'You are not authorized to view this chapter.'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ChapterContentSerializer(chapter)
        return Response(serializer.data)

    def put(self, request, book_id, chapter_id, format=None):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'You do not have permission to edit this chapter'}, status=status.HTTP_403_FORBIDDEN)

        serializer = ChapterContentSerializer(chapter, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def delete(self, request, book_id, chapter_id, format=None):
        try:
            chapter = Chapter.objects.get(pk=chapter_id, book_id=book_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        # Check if the user is authorized to delete the chapter
        if chapter.book.author != request.user:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        chapter.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


@api_view(['POST'])
def publish_action(request, book_id, chapter_id=None):
    action = request.data.get('action')
    user = request.user

    book = get_object_or_404(Book, id=book_id)

    # Проверяем, что пользователь является автором книги
    if user != book.author:
        return Response({'error': 'You are not authorized to change the publication status of this book.'},
                        status=status.HTTP_403_FORBIDDEN)

    if action == 'not_published':
        # Установить published = False для указанной главы и всех последующих глав
        if not chapter_id:
            return Response({'error': 'Chapter ID is required for this action.'},
                            status=status.HTTP_400_BAD_REQUEST)
        chapter = get_object_or_404(Chapter, id=chapter_id, book=book)
        chapter_number = chapter.chapter_number

        # Получаем все главы с номером >= текущей
        chapters_to_unpublish = Chapter.objects.filter(
            book=book,
            chapter_number__gte=chapter_number,
            published=True
        ).order_by('chapter_number')

        with transaction.atomic():
            # Массовое обновление без цикла
            chapters_to_unpublish.update(published=False)

        # Обновляем дату изменения книги
        book.updated = timezone.now()
        book.save()

        return Response({
            'message': f'Chapters from {chapter_number} onward have been unpublished successfully.'
        })

    elif action == 'publish_chapter':
        # Установить published = True для указанной главы
        if not chapter_id:
            return Response({'error': 'Chapter ID is required for this action.'},
                            status=status.HTTP_400_BAD_REQUEST)
        chapter = get_object_or_404(Chapter, id=chapter_id, book=book)

        # Проверяем, что все предыдущие главы опубликованы
        if Chapter.objects.filter(book=book, chapter_number__lt=chapter.chapter_number, published=False).exists():
            return Response({'error': 'You cannot publish this chapter until all previous chapters are published.'},
                            status=status.HTTP_403_FORBIDDEN)

        chapter.published = True
        chapter.save()

        # Отправляем уведомление
        send_book_update_notifications(book, chapter.title)
        return Response({'published': chapter.published, 'message': 'Chapter published successfully.'})

    elif action == 'publish_book':
        # Установить published = True для всех глав в книге
        unpublished_chapters = Chapter.objects.filter(book=book, published=False).order_by('chapter_number')

        with transaction.atomic():
            for chapter in unpublished_chapters:
                chapter.published = True
                chapter.save()
                # Опционально: отправлять уведомление для каждой главы
                send_book_update_notifications(book, chapter.title)

        return Response({'message': 'All chapters published successfully.'})

    else:
        return Response({'error': 'Invalid action.'}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
def add_author_note(request, book_id, chapter_id):
    try:
        chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
    except Chapter.DoesNotExist:
        return Response({'error': 'Chapter not found'}, status=404)

    # Data from request
    start = request.data.get('start')
    end = request.data.get('end')
    note_text = request.data.get('note_text')

    # Validate data
    if start is None or end is None or note_text is None:
        return Response({'error': 'Invalid request'}, status=400)

    # Create a new AuthorNote instance
    AuthorNote.objects.create(
        chapter=chapter,
        book=chapter.book,  # Set the book field
        author=request.user,
        start_position=start,
        end_position=end,
        note_text=note_text
    )

    return Response({'message': 'Note added successfully'})


@api_view(['GET'])
def get_author_notes(request, book_id, chapter_id):
    notes = AuthorNote.objects.filter(book_id=book_id, chapter_id=chapter_id, author=request.user)
    serialized_notes = AuthorNoteSerializer(notes, many=True)
    return Response(serialized_notes.data)


class BookNotesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id):
        notes = AuthorNote.objects.filter(book=book_id, author=request.user)
        serializer = AuthorNoteSerializer(notes, many=True)
        return Response(serializer.data)


class ChapterNotesView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, chapter_id):
        notes = AuthorNote.objects.filter(book_id=book_id, chapter_id=chapter_id, author=request.user)
        serializer = AuthorNoteSerializer(notes, many=True)
        return Response(serializer.data)


class ChapterDeleteView(APIView):
    permission_classes = [IsAuthenticated]

    def delete(self, request, book_id, chapter_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)

        if book.author != user:
            return Response({'error': 'You are not authorized to delete chapters from this book.'},
                            status=status.HTTP_403_FORBIDDEN)

        chapter = get_object_or_404(Chapter, id=chapter_id, book=book)

        chapter_count = Chapter.objects.filter(book=book).count()

        if chapter_count == 1:
            confirm_delete = request.query_params.get('confirm_delete', 'false').lower()
            if confirm_delete == 'true':
                chapter.delete()
                book.delete()
                return Response({'message': 'Chapter and book deleted successfully.'}, status=status.HTTP_200_OK)
            else:
                return Response({
                    'warning': 'You are about to delete the book, the changes are irreversible. Do You want to proceed?',
                    'book_will_be_deleted': True
                }, status=status.HTTP_200_OK)
        else:
            chapter.delete()
            return Response({'message': 'Chapter deleted successfully.'}, status=status.HTTP_200_OK)


class ChapterDownloadView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, chapter_id):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        fb2_content = create_fb2(chapter)
        response = HttpResponse(fb2_content, content_type='application/xml')
        response['Content-Disposition'] = f'attachment; filename="{chapter.title}.fb2"'
        return response


class ChapterUploadView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, chapter_id):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id)
        except Chapter.DoesNotExist:
            return Response({'error': 'Chapter not found'}, status=status.HTTP_404_NOT_FOUND)

        if chapter.book.author != request.user:
            return Response({'error': 'Permission denied'}, status=status.HTTP_403_FORBIDDEN)

        file = request.FILES.get('file')
        if not file or not file.name.endswith('.fb2'):
            return Response({'error': 'Invalid file format'}, status=status.HTTP_400_BAD_REQUEST)

        # Save file temporarily
        file_path = default_storage.save(f'tmp/{file.name}', file)
        with default_storage.open(file_path, 'rb') as fb2_file:
            content = parse_fb2(fb2_file)  # Parse FB2 file to extract content

        chapter.content = content
        chapter.save()

        # Cleanup the temporary file
        default_storage.delete(file_path)

        return Response({'message': 'Chapter uploaded successfully'})


class DownloadBookView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, format):
        book = get_object_or_404(Book, pk=book_id)

        if not book.can_user_download(request.user):
            return Response({'error': 'You are not allowed to download this book.'}, status=status.HTTP_403_FORBIDDEN)

        try:
            # Assuming you have a method to get the file path based on the format
            file_path = book.get_file_path(format)
            return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=f'{book.name}.{format}')
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_404_NOT_FOUND)


class SearchApiView(APIView):
    def get(self, request):
        query = request.query_params.get('q', '')
        books = Book.objects.filter(name__icontains=query)
        book_serializer = BookSerializer(books, many=True)

        # Filter authors who have at least one published book
        authors = User.objects.annotate(
            has_books=Exists(Book.objects.filter(author=OuterRef('pk')))
        ).filter(has_books=True, username__icontains=query).distinct()
        author_serializer = AuthorSerializer(authors, many=True)

        return Response({
            'books': book_serializer.data,
            'authors': author_serializer.data
        })


class BookSettingsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, book_id, format=None):
        book = get_object_or_404(Book, pk=book_id)

        if book.author != request.user:
            return Response({'error': 'You do not have permission to view the settings of this book'}, status=status.HTTP_403_FORBIDDEN)

        # Create the context for the serializer
        followers_usernames = request.user.follower_users.values_list('follower__username', flat=True)
        context = {'co_author_queryset': User.objects.filter(username__in=followers_usernames)}

        serializer = BookSettingsSerializer(book, context=context)
        return Response(serializer.data)

    def patch(self, request, book_id, format=None):
        book = get_object_or_404(Book, pk=book_id)

        if book.author != request.user:
            return Response({'error': 'You do not have permission to edit this book'}, status=status.HTTP_403_FORBIDDEN)

        # Create the context for the serializer
        followers_usernames = request.user.follower_users.values_list('follower__username', flat=True)
        context = {'co_author_queryset': User.objects.filter(username__in=followers_usernames)}

        serializer = BookSettingsSerializer(book, data=request.data, partial=True, context=context)
        if serializer.is_valid():
            if 'is_adult' in serializer.validated_data and serializer.validated_data['is_adult'] is True:
                if not serializer.validated_data.get('confirm_adult_content', False):
                    return Response({'error': 'Please confirm that you want to set this book as adult content.'},
                                    status=status.HTTP_400_BAD_REQUEST)
            serializer.save()
            return Response(serializer.data)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, book_id, format=None):
        action = request.data.get('action')
        if action == 'reset_settings':
            book = get_object_or_404(Book, pk=book_id)
            if book.author != request.user:
                return Response({'error': 'You do not have permission to edit this book'},
                                status=status.HTTP_403_FORBIDDEN)

            # Сбрасываем настройки книги
            book.name = "Book Name"
            book.co_author = None
            book.co_author2 = None
            book.genre = Genre.objects.get_or_create(name='Undefined')[0]
            book.subgenres.clear()
            book.description = "Book's Description"
            book.authors_note = "Author's Note"
            book.visibility = 'private'
            book.comment_access = 'public'
            book.download_access = 'public'
            book.demo_version = False
            book.save()

            return Response({'status': 'Book settings have been reset to default values.'}, status=status.HTTP_200_OK)

        # Обработка других действий или возвращение ошибки
        return Response({'error': 'Invalid action.'}, status=status.HTTP_400_BAD_REQUEST)


class IllustrationView(APIView):
    def get(self, request, *args, **kwargs):
        illustrations = Illustration.objects.filter(book__id=kwargs['book_id'])
        serializer = IllustrationSerializer(illustrations, many=True, context={'request': request})
        return Response(serializer.data)


class BookSaleView(APIView):
    pass


class BookTextAPIView(APIView):
    """
    Add text to a book.
    """

    def post(self, request, *args, **kwargs):
        book_id = request.session.get('book_id', None)

        if not book_id:
            return Response({"error": "Book ID not found in session."}, status=status.HTTP_400_BAD_REQUEST)

        book = get_object_or_404(Book, id=book_id)
        serializer = ChapterSerializers(data=request.data)

        if serializer.is_valid():
            chapter = serializer.save(book=book)
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@method_decorator(staff_member_required, name='dispatch')
class BooksUpdate(UpdateView):
    template_name = 'store/book_create.html'
    form_class = BooksForm
    success_url = '/'
    queryset = Book.objects.all()


@method_decorator(staff_member_required, name='dispatch')
class BooksDelete(DeleteView):
    template_name = 'store/book_delete.html'
    queryset = Book.objects.all()
    success_url = '/'


class SeriesCreateView(LoginRequiredMixin, CreateView):
    model = Series
    form_class = SeriesForm
    template_name = 'store/create_series.html'
    success_url = '/'

    def form_valid(self, form):
        series = form.save(commit=False)
        series.author = self.request.user
        series.save()
        return super().form_valid(form)


class SeriesDetailView(DetailView):
    model = Series
    context_object_name = 'series'
    template_name = 'store/series_detail.html'


class SeriesUpdateView(LoginRequiredMixin, UpdateView):
    model = Series
    form_class = SeriesForm
    template_name = 'store/update_series.html'

    def form_valid(self, form):
        series = form.save(commit=False)
        series.author = self.request.user
        series.save()
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('series_detail', args=[self.object.id])


class CommentListCreateView(APIView):
    def get_permissions(self):
        """
        Instantiates and returns the list of permissions that this view requires.
        """
        if self.request.method == 'POST':
            self.permission_classes = [IsAuthenticated, ]  # Only authenticated users can post comments
        else:
            self.permission_classes = [AllowAny, ]  # Allow anyone to read comments
        return super(CommentListCreateView, self).get_permissions()

    def get(self, request, book_id):
        book = get_object_or_404(Book, pk=book_id)
        comments = Comment.objects.filter(book=book, parent_comment__isnull=True).order_by('-rating')
        serialized_comments = CommentSerializer(comments, many=True, context={'request': request})
        return Response({'comments': serialized_comments.data})

    def post(self, request, book_id):
        book = get_object_or_404(Book, pk=book_id)
        if not book.can_user_comment(request.user):
            return Response({'error': 'You are not allowed to comment on this book.'}, status=status.HTTP_403_FORBIDDEN)

        data = request.data.copy()
        data['book'] = book.id
        serializer = CreateCommentSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            new_comment = serializer.save(user=request.user)
            # Вебсокет но для него требуется сервер отдельный редис с отдельной настройкой
            '''channel_layer = get_channel_layer()
            group_name = f'comments_{book_id}'
            async_to_sync(channel_layer.group_send)(
                group_name,
                {
                    'type': 'new_comment',
                    'message': StudioCommentSerializer(new_comment, context={'request': request}).data
                }
            )'''
            return Response(StudioCommentSerializer(new_comment, context={'request': request}).data,
                            status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class LikeCommentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, comment_id):

        comment = get_object_or_404(Comment, id=comment_id, book_id=book_id)
        if comment.user == request.user:
            return Response({'error': 'You cannot like your own comment.'}, status=status.HTTP_403_FORBIDDEN)

        CommentDislike.objects.filter(comment=comment, user=request.user).delete()
        like, created = CommentLike.objects.get_or_create(comment=comment, user=request.user)

        return Response({'status': 'liked' if created else 'like exists'}, status=status.HTTP_200_OK)


class DislikeCommentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, comment_id):

        comment = get_object_or_404(Comment, id=comment_id, book_id=book_id)
        if comment.user == request.user:
            return Response({'error': 'You cannot dislike your own comment.'}, status=status.HTTP_403_FORBIDDEN)

        CommentLike.objects.filter(comment=comment, user=request.user).delete()
        dislike, created = CommentDislike.objects.get_or_create(comment=comment, user=request.user)

        return Response({'status': 'disliked' if created else 'dislike exists'}, status=status.HTTP_200_OK)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_comment(request, comment_id):
    try:
        comment = Comment.objects.get(pk=comment_id, user=request.user)
    except Comment.DoesNotExist:
        return Response({'error': 'Comment not found or not owned by user.'}, status=status.HTTP_404_NOT_FOUND)

    comment.delete()
    return Response({'message': 'Comment deleted successfully.'}, status=status.HTTP_204_NO_CONTENT)


class ReviewListView(generics.ListCreateAPIView):
    serializer_class = ReviewSerializer

    def get_queryset(self):
        book_id = self.kwargs.get('book_id')
        return Review.objects.filter(book_id=book_id)

    def get_serializer_class(self):
        if self.request.method == 'POST':
            return ReviewCreateSerializer
        return ReviewSerializer

    def perform_create(self, serializer):
        serializer.save()


class LikeReviewView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id, review_id):
        review = get_object_or_404(Review, id=review_id, book_id=book_id)
        if review.author == request.user:
            return Response({'error': 'You cannot like your own review.'}, status=status.HTTP_403_FORBIDDEN)

        # Toggle the like
        if request.user in review.likes.all():
            review.likes.remove(request.user)
            status_message = 'like removed'
        else:
            review.likes.add(request.user)
            status_message = 'liked'

        return Response({'status': status_message}, status=status.HTTP_200_OK)



@api_view(['GET'])
def review_detail(request, review_id):
    review = get_object_or_404(Review, id=review_id)
    review.increase_views_count(request)

    serializer = ReviewSerializer(review)
    return Response(serializer.data)


def upvote_book(request, book_id):
    book = get_object_or_404(Book, id=book_id)

    # Check if the user has already downvoted the book
    if BookDownvote.objects.filter(user=request.user, book=book).exists():
        # User has previously downvoted the book, so remove the downvote
        BookDownvote.objects.filter(user=request.user, book=book).delete()

    # Perform the logic for handling the upvote action
    BookUpvote.objects.get_or_create(user=request.user, book=book)

    # Update the book rating
    book.rating = book.upvotes.count() - book.downvotes.count()
    book.save()

    # Redirect back to the same page
    return redirect(reverse('book_detail', kwargs={'pk': book.pk}))


def downvote_book(request, book_id):
    book = get_object_or_404(Book, id=book_id)

    # Check if the user has already upvoted the book
    if BookUpvote.objects.filter(user=request.user, book=book).exists():
        # User has previously upvoted the book, so remove the upvote
        BookUpvote.objects.filter(user=request.user, book=book).delete()

    # Perform the logic for handling the downvote action
    BookDownvote.objects.get_or_create(user=request.user, book=book)

    # Update the book rating
    book.rating = book.upvotes.count() - book.downvotes.count()
    book.save()

    # Redirect back to the same page
    return redirect(reverse('book_detail', kwargs={'pk': book.pk}))


class AddToLikedView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)
        library, _ = Library.objects.get_or_create(user=user)

        # Check if the book is adult and the user is under 18
        if book.is_adult:
            if not user.profile.date_of_birth:
                return Response({'error': 'Your age is not specified.'}, status=status.HTTP_400_BAD_REQUEST)
            age = (date.today() - user.profile.date_of_birth).days // 365
            if age < 18:
                return Response({'error': 'This book is for adults only.'}, status=status.HTTP_403_FORBIDDEN)

        # Check if the book is already liked
        if book in library.liked_books.all():
            return Response({'message': 'Book is already liked'}, status=status.HTTP_400_BAD_REQUEST)

        # Add book to the liked category
        library.liked_books.add(book)
        return Response({'message': 'Book added to liked successfully'}, status=status.HTTP_200_OK)


class Reader(APIView):
    permission_classes = [IsAuthenticatedOrReadOnly]

    def get(self, request, book_id):
        try:
            book = Book.objects.get(pk=book_id)
            # Add a check for the book's adult content flag
            if book.is_adult:
                if not request.user.is_authenticated:
                    return Response({'detail': 'Sorry, you need to log in and be older than 18 to read this book.'}, status=status.HTTP_401_UNAUTHORIZED)
                if not self.is_user_adult(request.user):
                    return Response({'detail': 'Sorry, this content is restricted to users over 18.'}, status=status.HTTP_403_FORBIDDEN)

        except Book.DoesNotExist:
            return Response({'detail': 'Book not found.'}, status=status.HTTP_404_NOT_FOUND)

        chapters = Chapter.objects.filter(book=book, published=True)
        user_is_author = request.user == book.author
        user_has_purchased = False
        if request.user.is_authenticated:
            user_has_purchased = book in request.user.library.purchased_books.all()

        can_access_all_chapters = user_is_author or user_has_purchased

        serialized_chapters = self.serialize_chapters(chapters, book, request.user, can_access_all_chapters)

        return Response(serialized_chapters)

    def is_user_adult(self, user):
        try:
            web_page_settings = WebPageSettings.objects.get(profile__user=user)
            dob = web_page_settings.date_of_birth
            if dob is None:
                # If DOB is not set, consider as not adult
                return False
            today = timezone.now().date()
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            return age >= 18
        except WebPageSettings.DoesNotExist:
            # Handle cases where the user does not have WebPageSettings or it's not configured properly
            return False

    def serialize_chapters(self, chapters, book, user, user_has_purchased):
        serialized_chapters = []
        for chapter in chapters:
            chapter_data = {
                'id': chapter.id,
                'title': chapter.title,
            }
            if chapter.is_free or user_has_purchased:
                chapter_data['content'] = chapter.content
            else:
                chapter_data['content'] = 'This content is locked. Please purchase the book to read.'
            serialized_chapters.append(chapter_data)

        return serialized_chapters


def is_user_adult(user):
    if not user.is_authenticated:
        return False
    try:
        dob = user.profile.webpagesettings.date_of_birth
        today = timezone.now().date()
        return (today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))) >= 18
    except (AttributeError, WebPageSettings.DoesNotExist):
        return False


class SingleChapterView(APIView):
    def get(self, request, book_id, chapter_id):
        try:
            chapter = Chapter.objects.get(book_id=book_id, id=chapter_id, published=True)
        except Chapter.DoesNotExist:
            return Response({'detail': 'Chapter not found or not published.'}, status=status.HTTP_404_NOT_FOUND)

        is_author = request.user.is_authenticated and chapter.book.author == request.user
        is_adult = is_user_adult(request.user)
        can_access = (
            (chapter.is_free and (not chapter.book.is_adult or is_adult)) or
            (is_book_purchased_by_user(chapter.book, request.user) and (not chapter.book.is_adult or is_adult)) or
            is_author
        )

        if not can_access:
            return Response({'detail': 'You do not have access to this chapter.'}, status=status.HTTP_403_FORBIDDEN)

        serialized_chapter = ChapterSerializers(chapter).data

        # Update user's book reading history if record_history is True
        if request.user.is_authenticated and request.user.profile.record_history:
            UserBookHistory.objects.update_or_create(
                user=request.user,
                book=chapter.book,
                defaults={'last_accessed': timezone.now()}
            )

        if not request.user.is_authenticated:
            history = get_unlogged_user_history(request)
            if len(history) >= 10:
                history.pop(0)  # Remove the oldest entry if history is full
            history.append(chapter.book.id)
            set_unlogged_user_history(request, history)

        return Response(serialized_chapter)


def get_unlogged_user_history(request):
    return request.session.get('unlogged_user_history', [])


def set_unlogged_user_history(request, book_ids):
    request.session['unlogged_user_history'] = book_ids
    request.session.modified = True  # Ensure the session is saved


class VoteBookView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)
        value = request.data.get('value')

        if str(value) not in ('1', '-1', '0'):
            return Response({'error': 'Invalid vote value.'}, status=status.HTTP_400_BAD_REQUEST)

        value = int(value)

        if value == 0:
            # Удаляем голос
            Vote.objects.filter(user=user, book=book).delete()
        else:
            # Обновляем или создаём голос
            vote, created = Vote.objects.update_or_create(
                user=user,
                book=book,
                defaults={'value': value}
            )

        # Обновляем рейтинг книги
        book.update_rating()

        return Response({'message': 'Vote recorded.', 'rating': book.rating}, status=status.HTTP_200_OK)


class PurchaseBookView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        profile = user.profile
        wallet = get_object_or_404(Wallet, profile=profile)

        book = get_object_or_404(Book, id=book_id)
        book_price = book.price

        if wallet.balance >= book_price:
            wallet.purchase(book, book_price)
            user.library.purchased_books.add(book)
            # Add additional logic if needed, like adding the book to the user's library
            return Response({'message': 'Book purchased successfully'})
        else:
            return Response({'error': 'Insufficient wallet balance'}, status=400)


class RefundBookView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, book_id):
        user = request.user
        book = get_object_or_404(Book, id=book_id)

        try:
            purchased_book = PurchasedBook.objects.get(library=user.library, book=book)
            # Check if the refund is within the allowed time frame (e.g., 30 days)
            if timezone.now() - purchased_book.purchase_date <= timedelta(days=30):
                user.library.purchased_books.remove(book)
                user.profile.wallet.balance += book.price
                user.profile.wallet.save()
                purchased_book.delete()
                return Response({'message': 'Book refunded successfully'})
            else:
                return Response({'error': 'Refund period has expired'}, status=400)
        except PurchasedBook.DoesNotExist:
            return Response({'error': 'Book not found in purchased books'}, status=404)


class HistoryView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        now = timezone.now()
        search_query = request.query_params.get('search', None)
        category_query = request.query_params.get('category', None)  # New query param for time category

        user_history = UserBookHistory.objects.filter(user=request.user).order_by('-last_accessed')

        if search_query:
            user_history = user_history.filter(book__name__icontains=search_query)

        # Define the time categories
        time_categories = {
            "Today": user_history.filter(last_accessed__date=now.date()),
            "Yesterday": user_history.filter(last_accessed__date=(now - timedelta(days=1)).date()),
            "Last Week": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=7), now.date() - timedelta(days=2)]),
            "Week Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=14), now.date() - timedelta(days=8)]),
            "Two Weeks Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=21), now.date() - timedelta(days=15)]),
            "Three Weeks Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=28), now.date() - timedelta(days=22)]),
            "Month Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=60), now.date() - timedelta(days=29)]),
            "Two Months Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=90), now.date() - timedelta(days=61)]),
            "Three Months Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=180), now.date() - timedelta(days=91)]),
            "Half Year Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=365), now.date() - timedelta(days=181)]),
            "A Year Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=730), now.date() - timedelta(days=366)]),
            "Two Years Ago": user_history.filter(
                last_accessed__date__range=[now.date() - timedelta(days=1095), now.date() - timedelta(days=731)]),
            "A Long Time Ago": user_history.filter(last_accessed__date__lte=now.date() - timedelta(days=1096)),
        }

        # Check if a specific category filter is applied
        if category_query and category_query in time_categories:
            filtered_history = time_categories[category_query]
            serialized_data = BookViewSerializer(filtered_history, many=True, context={'request': request}).data
            return Response({category_query: serialized_data})
        else:
            # Serialize each queryset if no specific filter is applied
            for time_category, queryset in time_categories.items():
                time_categories[time_category] = BookViewSerializer(queryset, many=True,
                                                                    context={'request': request}).data

            return Response(time_categories)


'''def record_history_view(request, book_id):
    if request.user.is_authenticated and request.user.profile.record_history:
        book = get_object_or_404(Book, id=book_id)
        update_user_book_history(request.user, book)
        return Response({'message': 'History recorded successfully'})
    return Response({'error': 'History recording is disabled or user is not authenticated'}, status=status.HTTP_403_FORBIDDEN)'''


def record_history_view(request, book_id): # Alter record history, надо потестить позже
    book = get_object_or_404(Book, id=book_id)

    # Всегда логируем доступ
    log_book_access(request.user, book, request)

    # Записываем в историю пользователя, только если это разрешено
    if request.user.is_authenticated and request.user.profile.record_history:
        update_user_book_history(request.user, book)
        return Response({'message': 'History recorded successfully'})

    return Response({'message': 'Access logged'})


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def delete_history(request):
    UserBookHistory.objects.filter(user=request.user).delete()
    return Response({'message': 'History deleted successfully'}, status=status.HTTP_200_OK)


logger = logging.getLogger('user_history')


def update_user_book_history(user, book):
    # Check if there's an existing history entry for this user and book
    history_entry, created = UserBookHistory.objects.get_or_create(user=user, book=book)

    # Update the last_accessed timestamp to the current time
    history_entry.last_accessed = timezone.now()
    history_entry.save()


def update_user_book_history_with_logging(user, book):
    # Update the user book history
    update_user_book_history(user, book)

    # Log the event
    logger.debug(f'Updated history for user {user.username} and book {book.name}')


@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def update_history_settings(request):
    if request.method == 'GET':
        return Response({'record_history': request.user.profile.record_history})

    if request.method == 'POST':
        # Изменяем текущую настройку
        request.user.profile.record_history = not request.user.profile.record_history
        request.user.profile.save()
        return Response({'record_history': request.user.profile.record_history})


class UnloggedUserHistoryView(APIView):
    def get(self, request):
        book_ids = get_unlogged_user_history(request)
        books = Book.objects.filter(id__in=book_ids).order_by('-id')  # Ensure the order reflects the most recent views
        serializer = BookSerializer(books, many=True)
        return Response(serializer.data)

    def post(self, request, book_id):
        book_ids = get_unlogged_user_history(request)

        # Attempt to add the book to history
        if book_id not in book_ids:
            book_ids.append(book_id)
            if len(book_ids) > 10:
                book_ids.pop(0)  # Remove the oldest entry if the limit is exceeded

        set_unlogged_user_history(request, book_ids)
        return Response({'message': 'Book added to history'})


class NewsNotificationsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        settings = user.notification_settings

        # Получаем настройки пользователя для категорий библиотеки
        library_categories = []
        if settings.library_reading_updates:
            library_categories.append('reading')
        if settings.library_wishlist_updates:
            library_categories.append('wish_list')
        if settings.library_liked_updates:
            library_categories.append('liked')
        if settings.library_favourite_updates:
            library_categories.append('favorites')

        # Получаем книги из библиотеки пользователя в указанных категориях
        books = Book.objects.none()
        for category in library_categories:
            books |= getattr(user.library, f'{category}_books').all()
        books = books.distinct()

        # Фильтруем уведомления по типу и книгам
        notifications = Notification.objects.filter(
            recipient=user.profile,
            notification_type='book_update',
            book__in=books
        ).order_by('-timestamp')

        # Группируем уведомления по книгам
        books_dict = {}
        for notification in notifications:
            book_id = notification.book_id
            if book_id not in books_dict:
                books_dict[book_id] = {
                    'book': notification.book,
                    'notifications': [],
                    'updates_count': 0
                }
            books_dict[book_id]['notifications'].append(notification)
            books_dict[book_id]['updates_count'] += 1

        # Подготавливаем данные для сериализации
        serialized_data = []
        for info in books_dict.values():
            # Сортируем уведомления по времени
            sorted_notifications = sorted(info['notifications'], key=lambda n: n.timestamp, reverse=True)

            # Формируем список обновлений
            updates_list = [
                {
                    'id': n.id,
                    'chapter_title': n.chapter_title if n.chapter_title else n.chapter.title if n.chapter else '',
                    'formatted_timestamp': n.timestamp.strftime('%Y-%m-%d %H:%M:%S')
                }
                for n in sorted_notifications
            ]

            # Добавляем книгу и информацию об обновлениях
            serialized_data.append({
                'book': NewsInfoSerializer(info['book'], context={'request': request}).data,
                'updates_count': info['updates_count'],
                'updates_list': updates_list
            })

        # Сортируем книги по дате последнего обновления
        serialized_data.sort(key=lambda x: x['updates_list'][0]['formatted_timestamp'], reverse=True)

        return Response(serialized_data)


def parse_fb2_and_create_chapters(file_path, book):
    try:
        print(f"Parsing FB2 file at: {file_path}")
        tree = etree.parse(file_path)
        root = tree.getroot()
        ns = {'fb': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
        print("Successfully parsed XML tree.")

        last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0

        sections = root.findall('.//fb:section', namespaces=ns)
        print(f"Found {len(sections)} sections in the FB2 file.")

        if not sections:
            print("No sections found in the FB2 file.")
            raise ValueError("No sections found in the FB2 file.")

        for section in sections:
            title_element = section.find('./fb:title', namespaces=ns)
            title = "".join(title_element.itertext()) if title_element is not None else "Untitled Chapter"

            content = ""
            # Ищем все параграфы внутри секции, но исключаем те, которые находятся внутри <title>
            paragraphs = section.xpath('./fb:section/*[not(self::fb:title)]//fb:p | ./fb:p', namespaces=ns)
            if not paragraphs:
                print(f"No paragraphs found in section titled '{title}'.")
                continue  # Пропускаем секцию без содержимого

            for p in paragraphs:
                content += "\n".join(p.itertext()) + "\n"

            next_chapter_number = last_chapter_number + 1
            last_chapter_number = next_chapter_number  # Update last chapter number for the next iteration

            Chapter.objects.create(
                book=book,
                title=title.strip(),
                content=content.strip(),
                chapter_number=next_chapter_number,
                created=timezone.now(),
                updated=timezone.now(),
                published=False
            )
            print(f"Created chapter {next_chapter_number}: {title}")

    except Exception as e:
        print(f"Error parsing FB2 file: {e}")
        import traceback
        traceback.print_exc()
        raise e  # Пробрасываем исключение для обработки в вызывающем коде


def parse_docx_and_create_chapters(file_path, book):
    doc = Document(file_path)
    chapters = []
    current_chapter = []
    chapter_title = None
    last_chapter_number = 0

    # Регулярное выражение для поиска заголовков глав на английском языке
    chapter_heading_patterns = [
        r'^chapter\s+\d+',          # "Chapter 1"
        r'^chapter\s+[ivx]+',       # "Chapter IX"
        r'^chapter\s+\w+',          # "Chapter One"
        r'^ch\.\s*\d+',             # "Ch. 1"
        r'^\d+\.\s',                # "1. Introduction"
        r'^\w+\s*$'                 # Single word titles
    ]

    for para in doc.paragraphs:
        para_text = para.text.strip()
        # Проверяем, является ли параграф заголовком главы
        if any(re.match(pattern, para_text, re.IGNORECASE) for pattern in chapter_heading_patterns):
            if current_chapter:
                chapters.append((chapter_title or f"Chapter {last_chapter_number}", '\n'.join(current_chapter)))
                current_chapter = []
            chapter_title = para_text
            last_chapter_number += 1
        else:
            current_chapter.append(para_text)

    if current_chapter:
        chapters.append((chapter_title or f"Chapter {last_chapter_number}", '\n'.join(current_chapter)))

    if not chapters:
        # Если главы не найдены, создаем одну главу с полным текстом
        full_text = '\n'.join(para.text for para in doc.paragraphs)
        Chapter.objects.create(
            book=book,
            title="Full Text",
            content=full_text.strip(),
            chapter_number=1,
            created=timezone.now(),
            updated=timezone.now(),
            published=False
        )
        print("No chapters found in the DOCX file. Created a single chapter with full content.")
    else:
        for i, (title, content) in enumerate(chapters, start=1):
            Chapter.objects.create(
                book=book,
                title=title or f"Chapter {i}",
                content=content.strip(),
                chapter_number=i,
                created=timezone.now(),
                updated=timezone.now(),
                published=False
            )
        print(f"Created {len(chapters)} chapters from DOCX file.")


def parse_pdf_and_create_chapters(file_path, book):
    doc = fitz.open(file_path)
    toc = doc.get_toc()

    if toc:
        # Используем оглавление PDF для разделения на главы
        chapters = []
        for idx, entry in enumerate(toc):
            level, title, page_num = entry
            start_page = page_num - 1  # Нумерация страниц начинается с 0
            if idx + 1 < len(toc):
                end_page = toc[idx + 1][2] - 2
            else:
                end_page = len(doc) - 1

            content = ""
            for page in range(start_page, end_page + 1):
                content += doc.load_page(page).get_text()

            chapters.append((title.strip(), content.strip()))

        # Создаем главы
        for i, (title, content) in enumerate(chapters, start=1):
            if content:
                Chapter.objects.create(
                    book=book,
                    title=title or f"Chapter {i}",
                    content=content,
                    chapter_number=i,
                    created=timezone.now(),
                    updated=timezone.now(),
                    published=False
                )
        print(f"Created {len(chapters)} chapters from PDF file using TOC.")
    else:
        # Если оглавление отсутствует, пытаемся разделить по заголовкам глав
        full_text = ""
        for page in doc:
            full_text += page.get_text()

        # Используем регулярное выражение для поиска заголовков глав
        chapter_heading_pattern = re.compile(
            r'(Chapter\s+\d+|CHAPTER\s+\d+|Chapter\s+[IVX]+|CHAPTER\s+[IVX]+|Chapter\s+\w+|CHAPTER\s+\w+)',
            re.IGNORECASE
        )

        matches = list(chapter_heading_pattern.finditer(full_text))

        if matches:
            chapters = []
            for idx, match in enumerate(matches):
                title = match.group().strip()
                start_idx = match.end()
                end_idx = matches[idx + 1].start() if idx + 1 < len(matches) else len(full_text)
                content = full_text[start_idx:end_idx].strip()
                if content:
                    chapters.append((title, content))

            # Создаем главы
            for i, (title, content) in enumerate(chapters, start=1):
                Chapter.objects.create(
                    book=book,
                    title=title,
                    content=content,
                    chapter_number=i,
                    created=timezone.now(),
                    updated=timezone.now(),
                    published=False
                )
            print(f"Created {len(chapters)} chapters from PDF file using regex.")
        else:
            # Если главы не найдены, создаем одну главу с полным текстом
            Chapter.objects.create(
                book=book,
                title="Full Text",
                content=full_text.strip(),
                chapter_number=1,
                created=timezone.now(),
                updated=timezone.now(),
                published=False
            )
            print("No chapters found in the PDF file. Created a single chapter with full content.")

    doc.close()


def parse_txt_and_create_chapters(file_path, book):
    import chardet
    # Определяем кодировку файла
    with open(file_path, 'rb') as f:
        raw_data = f.read()
        result = chardet.detect(raw_data)
        encoding = result['encoding']
        confidence = result['confidence']
        print(f"Detected file encoding: {encoding} with confidence {confidence}")

        # Если не удалось определить кодировку или уверенность низкая, используем кодировку по умолчанию
        if encoding is None or confidence < 0.8:
            encoding = 'utf-8'
            print("Using default encoding: utf-8")

    # Читаем файл с определённой кодировкой
    try:
        with open(file_path, 'r', encoding=encoding) as file:
            text = file.read()
    except UnicodeDecodeError:
        # Если ошибка декодирования, пытаемся с кодировкой 'latin-1'
        print(f"Failed to decode with encoding {encoding}, trying 'latin-1'")
        with open(file_path, 'r', encoding='latin-1') as file:
            text = file.read()

    # Разделяем текст на главы
    chapter_heading_pattern = re.compile(
        r'(Chapter\s+\d+|CHAPTER\s+\d+|Chapter\s+[IVX]+|CHAPTER\s+[IVX]+|Chapter\s+\w+|CHAPTER\s+\w+)',
        re.IGNORECASE
    )
    chapters = chapter_heading_pattern.split(text)

    if len(chapters) > 1:
        # Первой в списке может быть пролог или введение
        for i in range(1, len(chapters), 2):
            title = chapters[i].strip()
            content = chapters[i+1].strip() if i+1 < len(chapters) else ''
            Chapter.objects.create(
                book=book,
                title=title,
                content=content,
                chapter_number=(i+1)//2,
                created=timezone.now(),
                updated=timezone.now(),
                published=False
            )
        print(f"Created {(len(chapters)-1)//2} chapters from TXT file using regex.")
    else:
        # Если главы не найдены, создаем одну главу с полным текстом
        Chapter.objects.create(
            book=book,
            title="Full Text",
            content=text.strip(),
            chapter_number=1,
            created=timezone.now(),
            updated=timezone.now(),
            published=False
        )
        print("No chapters found in the TXT file. Created a single chapter with full content.")


def parse_epub_and_create_chapters(file_path, book):
    import os
    import zipfile
    from bs4 import BeautifulSoup
    from django.utils import timezone
    from django.db.models import Max
    from .models import Chapter

    with zipfile.ZipFile(file_path, 'r') as zf:
        # Читаем container.xml, чтобы найти путь к файлу OPF
        with zf.open('META-INF/container.xml') as container_file:
            container_soup = BeautifulSoup(container_file, 'xml')
            rootfile = container_soup.find('rootfile')
            opf_path = rootfile['full-path']

        # Открываем файл OPF
        with zf.open(opf_path) as opf_file:
            opf_soup = BeautifulSoup(opf_file, 'xml')
            manifest_items = opf_soup.find_all('item')
            manifest = {item['id']: item['href'] for item in manifest_items}

            # Ищем элемент TOC
            toc_item = opf_soup.find('item', {'properties': 'nav'})
            if toc_item:
                toc_href = toc_item['href']
            else:
                toc_item = opf_soup.find('item', {'media-type': 'application/x-dtbncx+xml'})
                if toc_item:
                    toc_href = toc_item['href']
                else:
                    print("TOC not found in EPUB file.")
                    return

        # Вычисляем путь к файлу TOC
        opf_dir = os.path.dirname(opf_path)
        toc_file_path = os.path.normpath(os.path.join(opf_dir, toc_href))

        # Открываем файл TOC
        with zf.open(toc_file_path) as toc_file:
            toc_soup = BeautifulSoup(toc_file, 'xml')

            # Обрабатываем navPoints
            nav_points = toc_soup.find_all('navPoint')
            if not nav_points:
                print("No navPoints found in TOC.")
                return

            chapters_info = []
            for nav_point in nav_points:
                nav_label = nav_point.find('text').text.strip()
                content_src = nav_point.content['src']
                chapters_info.append({'title': nav_label, 'src': content_src})

        last_chapter_number = book.chapters.aggregate(Max('chapter_number'))['chapter_number__max'] or 0

        for chapter_info in chapters_info:
            # Обработка пути к файлу содержимого
            content_src = chapter_info['src'].split('#')[0]
            content_file_path = os.path.normpath(os.path.join(opf_dir, content_src))

            with zf.open(content_file_path) as file:
                soup = BeautifulSoup(file, 'html.parser')

                # Удаляем заголовок из содержимого главы
                if soup.body:
                    # Предполагаем, что заголовок находится в первом <h1>, <h2> или <h3>
                    for tag_name in ['h1', 'h2', 'h3']:
                        header = soup.body.find(tag_name)
                        if header and header.get_text(strip=True).strip().lower() == chapter_info['title'].lower():
                            header.decompose()
                            break

                    # Собираем содержимое главы с сохранением абзацев
                    content_parts = []
                    for element in soup.body.descendants:
                        if isinstance(element, str):
                            continue  # Пропускаем текстовые узлы вне тегов
                        if element.name in ['p', 'div', 'h1', 'h2', 'h3', 'blockquote']:
                            text = element.get_text(strip=True)
                            if text:
                                content_parts.append(text)
                    content = '<br/>'.join(content_parts)
                else:
                    content = ""

                next_chapter_number = last_chapter_number + 1
                last_chapter_number = next_chapter_number

                Chapter.objects.create(
                    book=book,
                    title=chapter_info['title'],
                    content=content,
                    chapter_number=next_chapter_number,
                    created=timezone.now(),
                    updated=timezone.now(),
                    published=False
                )
                print(f"Created chapter {next_chapter_number}: {chapter_info['title']}")

    # Если ни один файл не был обработан
    if not chapters_info:
        print("No chapters found in EPUB file.")


def get_file_type_by_magic(file_path):
    mime = magic.Magic(mime=True)
    mime_type = mime.from_file(file_path)
    return mime_type


class BookFileUploadView(APIView):
    handlers = {
        'application/epub+zip': parse_epub_and_create_chapters,
        'application/pdf': parse_pdf_and_create_chapters,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': parse_docx_and_create_chapters,
        'text/plain': parse_txt_and_create_chapters,
        'application/x-fictionbook+xml': parse_fb2_and_create_chapters,
        'application/xml': parse_fb2_and_create_chapters,  # Добавьте этот MIME-тип
        'text/xml': parse_fb2_and_create_chapters,  # Добавьте этот MIME-тип
        'fb2': parse_fb2_and_create_chapters,
        # ... другие MIME-типы ...
    }

    def post(self, request):
        serializer = BookFileSerializer(data=request.data)
        if serializer.is_valid():
            book = Book.objects.create(author=request.user, book_type='epic_novel')
            book_file = serializer.save(book=book)

            file_path = book_file.file.path
            print(f"Uploaded file path: {file_path}")
            extension = os.path.splitext(file_path)[1].lower()
            print(f"File extension: {extension}")


            # Используем python-magic для определения MIME-типа
            file_type = get_file_type_by_magic(book_file.file.path)
            print(f"Guessed file type using magic: {file_type}")

            if file_type is None:
                book.delete()
                book_file.file.delete()
                return Response({"error": "Unsupported file type"}, status=status.HTTP_400_BAD_REQUEST)

            handler = self.handlers.get(file_type)
            if handler:
                try:
                    handler(book_file.file.path, book)

                    # Получаем ID первой главы
                    first_chapter = book.chapters.order_by('chapter_number').first()
                    if first_chapter:
                        first_chapter_id = first_chapter.id
                    else:
                        # Если главы не созданы, возвращаем ошибку
                        book.delete()
                        book_file.file.delete()
                        return Response({"error": "Failed to create chapters"},
                                        status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                    # Создаем простой ответ без использования дополнительных сериализаторов
                    response_data = {
                        "message": "Book uploaded and processed successfully",
                        "book_id": book.id,
                        "first_chapter_id": first_chapter_id
                    }
                    return Response(response_data, status=status.HTTP_201_CREATED)
                except Exception as e:
                    print(f"Error in handler: {e}")
                    traceback.print_exc()
                    # Удаляем книгу и файл в случае ошибки
                    book.delete()
                    book_file.file.delete()
                    return Response({"error": f"Error processing file: {str(e)}"},
                                    status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            else:
                book.delete()
                book_file.file.delete()
                return Response({"error": "No handler for file type"}, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)




